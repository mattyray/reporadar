"""Tests for resume parsing task and views."""

import json
from unittest.mock import MagicMock, patch

import pytest
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.utils import timezone

from apps.resumes.models import ResumeProfile
from apps.resumes.tasks import _build_parse_prompt, _extract_text, parse_resume

User = get_user_model()


@pytest.fixture
def user(db):
    return User.objects.create_user(
        username="testuser", email="test@example.com", password="testpass123"
    )


@pytest.fixture
def resume_profile(user, tmp_path):
    """Create a ResumeProfile with a fake PDF file."""
    fake_pdf = tmp_path / "resume.pdf"
    fake_pdf.write_text("fake pdf content for testing")

    profile = ResumeProfile.objects.create(
        user=user,
        original_file="resumes/test_resume.pdf",
        file_type="pdf",
    )
    return profile


# --- Unit tests for pure functions ---


def test_build_parse_prompt_contains_resume_text():
    prompt = _build_parse_prompt("John Doe, 5 years Django experience")
    assert "John Doe" in prompt
    assert "tech_stack" in prompt
    assert "key_projects" in prompt
    assert "story_hook" in prompt
    assert "JSON" in prompt


def test_build_parse_prompt_requests_no_markdown():
    prompt = _build_parse_prompt("some resume text")
    assert "no markdown code fences" in prompt.lower()


# --- Tests for parse_resume task ---


@pytest.mark.django_db
@patch("apps.resumes.tasks._extract_text")
@patch("apps.resumes.tasks.anthropic")
def test_parse_resume_updates_profile(mock_anthropic, mock_extract, resume_profile):
    mock_extract.return_value = "John Doe, Django developer, 5 years exp"

    parsed_response = {
        "summary": "Full-stack Django developer with 5 years experience.",
        "tech_stack": ["Django", "Python", "PostgreSQL", "React"],
        "key_projects": [
            {"name": "RepoRadar", "description": "GitHub prospect finder", "tech": ["Django", "React"]}
        ],
        "years_experience": 5,
        "story_hook": "Self-taught developer who built a SaaS product from scratch.",
    }

    mock_message = MagicMock()
    mock_message.content = [MagicMock(text=json.dumps(parsed_response))]
    mock_client = MagicMock()
    mock_client.messages.create.return_value = mock_message
    mock_anthropic.Anthropic.return_value = mock_client

    with patch.dict("os.environ", {"ANTHROPIC_API_KEY": "test-key"}):
        result = parse_resume(resume_profile.id)

    resume_profile.refresh_from_db()
    assert resume_profile.summary == "Full-stack Django developer with 5 years experience."
    assert resume_profile.tech_stack == ["Django", "Python", "PostgreSQL", "React"]
    assert resume_profile.years_experience == 5
    assert resume_profile.story_hook == "Self-taught developer who built a SaaS product from scratch."
    assert len(resume_profile.key_projects) == 1
    assert resume_profile.parsed_at is not None
    assert result["status"] == "parsed"


@pytest.mark.django_db
@patch("apps.resumes.tasks._extract_text")
def test_parse_resume_fails_without_api_key(mock_extract, resume_profile):
    mock_extract.return_value = "some text"

    with patch.dict("os.environ", {}, clear=True):
        with pytest.raises(RuntimeError, match="ANTHROPIC_API_KEY"):
            parse_resume(resume_profile.id)


@pytest.mark.django_db
@patch("apps.resumes.tasks._extract_text")
def test_parse_resume_fails_on_empty_text(mock_extract, resume_profile):
    mock_extract.return_value = None

    with patch.dict("os.environ", {"ANTHROPIC_API_KEY": "test-key"}):
        with pytest.raises(RuntimeError, match="Could not extract text"):
            parse_resume(resume_profile.id)


@pytest.mark.django_db
@patch("apps.resumes.tasks._extract_text")
@patch("apps.resumes.tasks.anthropic")
def test_parse_resume_fails_on_invalid_json(mock_anthropic, mock_extract, resume_profile):
    mock_extract.return_value = "some resume text"

    mock_message = MagicMock()
    mock_message.content = [MagicMock(text="not valid json at all")]
    mock_client = MagicMock()
    mock_client.messages.create.return_value = mock_message
    mock_anthropic.Anthropic.return_value = mock_client

    with patch.dict("os.environ", {"ANTHROPIC_API_KEY": "test-key"}):
        with pytest.raises(RuntimeError, match="invalid JSON"):
            parse_resume(resume_profile.id)


# --- Tests for extract_text helpers ---


def test_extract_docx_text(tmp_path):
    """Test DOCX extraction with python-docx."""
    try:
        import docx

        doc = docx.Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Django Developer")
        path = tmp_path / "test.docx"
        doc.save(str(path))

        profile = MagicMock()
        profile.file_type = "docx"
        profile.original_file.path = str(path)

        text = _extract_text(profile)
        assert "John Doe" in text
        assert "Django Developer" in text
    except ImportError:
        pytest.skip("python-docx not installed")
